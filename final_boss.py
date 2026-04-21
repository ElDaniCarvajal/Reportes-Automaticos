#!/usr/bin/env python3
# -*- coding: utf-8 -*-

from __future__ import annotations

import argparse
import os
import re
import shutil
import subprocess
import sys
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Dict, List, Optional, Sequence, Tuple


# ============================================================
# CONFIGURACIÓN EDITABLE
# ============================================================
BASE_DIR = Path(__file__).resolve().parent

AMP_SCRIPT = BASE_DIR / "amp_unificado.py"
TENABLE_SCRIPT = BASE_DIR / "Tenable-images-final.py"
TEMPLATE_DOCX = BASE_DIR / "Argentina_servicios de ciberseguridad_febrero 2026.docx"

BASE_OUTPUT_DIR = BASE_DIR / "reportes_finales_docx"
CURRENT_MONTH_FOLDER = datetime.now().strftime("%Y-%m")
OUTPUT_DIR = BASE_OUTPUT_DIR / CURRENT_MONTH_FOLDER

REPORT_PERIOD = ""

TEMPLATE_ENTITY_NAME = "Argentina"
TEMPLATE_PERIOD_TEXT: Optional[str] = None
REPLACE_TEXT_INSIDE_DOCX = False

# Slots internos del DOCX
DOCX_IMAGE_MAP: Dict[str, Tuple[str, str]] = {
    # AMP
    "word/media/image4.png": ("amp", "compromises.png"),
    "word/media/image5.png": ("amp", "threats.png"),
    "word/media/image6.png": ("amp", "devices.png"),

    # Tenable
    "word/media/image8.png": ("tenable", "03_vulns_severity.png"),
    "word/media/image9.png": ("tenable", "04_vulns_top10.png"),
    "word/media/image10.png": ("tenable", "05_vulns_spotlight.png"),
    "word/media/image11.png": ("tenable", "identity_exposure_donuts_4k.png"),
    "word/media/image12.png": ("tenable", "01_cis_controls_by_asset.png"),
    "word/media/image13.png": ("tenable", "02_cis_controls_totals.png"),
    "word/media/image14.png": ("tenable", "01_audits_list_4k.png"),
    "word/media/image15.png": ("tenable", "02_audit_detail_4k.png"),
}


@dataclass(frozen=True)
class Target:
    display_name: str
    amp_name: str
    tenable_category: str
    tenable_value: str
    output_filename: Optional[str] = None


# ============================================================
# TARGETS EXACTOS
# ============================================================
TARGETS: Tuple[Target, ...] = (
    # Países
    Target("Argentina", "ARGENTINA", "Omnilife", "Argentina"),
    Target("Bolivia", "BOLIVIA", "Omnilife", "Bolivia"),
    Target("Brasil", "BRASIL", "Omnilife", "Brasil"),
    Target("Chile", "CHILE", "Omnilife", "Chile"),
    Target("Colombia", "COLOMBIA", "Omnilife", "Colombia-Todos"),
    Target("Costa Rica", "COSTA RICA", "Omnilife", "Costa Rica"),
    Target("Ecuador", "ECUADOR", "Omnilife", "Ecuador"),
    Target("El Salvador", "EL SALVADOR", "Omnilife", "El Salvador"),
    Target("España", "ESPAÑA", "Omnilife", "España"),
    Target("Estados Unidos", "USA", "Omnilife", "Estados Unidos"),
    Target("Guatemala", "GUATEMALA", "Omnilife", "Guatemala"),
    Target("Nicaragua", "NICARAGUA", "Omnilife", "Nicaragua"),
    Target("Panamá", "PANAMA", "Omnilife", "Panamá"),
    Target("Paraguay", "PARAGUAY", "Omnilife", "Paraguay"),
    Target("Perú", "PERU", "Omnilife", "Perú"),
    Target(
        "República Dominicana",
        "Republica Dominicana",
        "Omnilife",
        "Republica Dominicana",
        output_filename="Republica Dominicana_servicios de ciberseguridad",
    ),
    Target("Rusia", "RUSIA", "Omnilife", "Rusia"),
    Target("Uruguay", "URUGUAY", "Omnilife", "Uruguay"),

    # Razones sociales
    Target("Arte y Cultura", "Arte y Cultura Omnilife A.C", "Razón Social", "Arte y Cultura (Museo JV)"),
    Target("Consorcio VAV", "Consorcio Vav, S.A. de C.V", "Razón Social", "Consorcio VAV"),
    Target("Educare", "Educare", "Razón Social", "Educare"),
    Target(
        "Fundación Jorge Vergara",
        "Fundación Jorge vergara, A.C",
        "Razón Social",
        "Fundación Jorge Vergara",
        output_filename="Fundacion Jorge Vergara_servicios de ciberseguridad",
    ),
    Target("OML Seguros", "OML Agente de Seguros y de Fianzas, S.A. de C.V", "Razón Social", "OML Seguros"),
    Target("Omnia de Guadalajara", "Omnia de Guadalajara, S.A. de C.V", "Razón Social", "Omnia de Guadalajara"),
    Target("Omnidata Internacional", "Omnidata Internacional", "Razón Social", "Omnidata Internacional"),
    Target("Omnihumana", "Omnihumana S.A. de C.V", "Razón Social", "Omnihumana"),
    Target("Omnilife de México", "Omnilife de México S.A de C.V", "Razón Social", "Omnilife de México"),
    Target("Omnilife Manufactura", "Omnilife Manufactura", "Razón Social", "Omnilife Manufactura"),
    Target("Omnipromotora", "Omni promotora de Negocios SA", "Razón Social", "Omnipromotora"),
    Target("Omnisky", "Omnisky SA de CV", "Razón Social", "Omnisky"),
    Target("Planeta Morado", "Planeta Morado , A.C", "Razón Social", "Planeta Morado"),
    Target(
        "Seytú",
        "Seytu Cosmética, S.A. de C.V",
        "Razón Social",
        "Seytú",
        output_filename="Seytu_servicios de ciberseguridad",
    ),
    Target("Templo Mayor", "Templo Mayor de Chivas, A.C", "Razón Social", "Templo Mayor"),
    Target("Transur", "Transur, S.A. de C.V", "Razón Social", "Transur"),
)


# ============================================================
# HELPERS
# ============================================================
def norm(s: str) -> str:
    s = (s or "").strip().lower()
    replacements = {
        "á": "a", "é": "e", "í": "i", "ó": "o", "ú": "u", "ü": "u", "ñ": "n",
    }
    for a, b in replacements.items():
        s = s.replace(a, b)
    s = re.sub(r"[^a-z0-9]+", " ", s)
    return re.sub(r"\s+", " ", s).strip()


def safe_name(s: str) -> str:
    s = re.sub(r'[<>:"/\\|?*\x00-\x1F]', "_", (s or "").strip())
    s = re.sub(r"\s+", " ", s).strip().rstrip(". ")
    return s or "unnamed"


def output_filename_for(target: Target) -> str:
    base = target.output_filename or f"{target.display_name}_servicios de ciberseguridad"
    if REPORT_PERIOD.strip():
        return f"{safe_name(base)}_{REPORT_PERIOD}.docx"
    return f"{safe_name(base)}.docx"


def require_exists(path: Path, label: str) -> None:
    if not path.exists():
        raise FileNotFoundError(f"No existe {label}: {path}")


def run_python_script(script: Path, workdir: Path, extra_env: Optional[Dict[str, str]] = None) -> None:
    env = os.environ.copy()
    if extra_env:
        env.update(extra_env)

    cmd = [sys.executable, str(script)]
    print(f"[INFO] Ejecutando: {' '.join(cmd)}")
    subprocess.run(cmd, cwd=str(workdir), env=env, check=True)


def find_latest_timestamp_dir(root: Path) -> Path:
    """
    Para AMP sí existe normalmente:
      amp_reportes_unificado/unificado_YYYYMMDD_HHMMSS

    Para Tenable NO debemos usar esta lógica sobre la raíz final.
    """
    if not root.is_dir():
        return root

    children = [p for p in root.iterdir() if p.is_dir()]
    if not children:
        return root

    timestamp_like = []
    for child in children:
        name = child.name.lower().strip()
        if name.startswith("unificado_") or re.fullmatch(r"\d{4}-\d{2}", name):
            timestamp_like.append(child)

    if timestamp_like:
        return sorted(timestamp_like, key=lambda p: p.stat().st_mtime, reverse=True)[0]

    return root


def list_amp_group_dirs(amp_root: Path) -> Dict[str, Path]:
    """
    AMP suele quedar así:
      amp_reportes_unificado/
        unificado_20260320_084326/
          ARGENTINA/
          BOLIVIA/
          ...
    """
    out: Dict[str, Path] = {}
    if not amp_root.is_dir():
        return out

    real_root = find_latest_timestamp_dir(amp_root)

    for child in real_root.iterdir():
        if child.is_dir():
            out[norm(child.name)] = child

    return out


def looks_like_tenable_category_dir(path: Path) -> bool:
    if not path.is_dir():
        return False
    name = norm(path.name)
    return name in {
        "asm",
        "chivas",
        "educare",
        "manufactura mexico",
        "omnilife",
        "omnilife de mexico",
        "razon social",
        "un",
    }


def resolve_tenable_base(tenable_root: Path) -> Path:
    """
    Tenable, según tu salida real, queda así:
      tenable_reportes/2026-03/
        ASM/
        Chivas/
        EDUCARE/
        Manufactura-Mexico/
        Omnilife/
        Omnilife de Mexico/
        Razón Social/
        UN/
        identity_exposure_donuts_4k.png
        index.json
        ...

    O en temporales:
      tenable_reportes_unificados/
        ASM/
        Chivas/
        ...

    Aquí NO debemos bajar a una subcarpeta "más reciente" al azar.
    """
    if not tenable_root.is_dir():
        return tenable_root

    if any(looks_like_tenable_category_dir(p) for p in tenable_root.iterdir()):
        return tenable_root

    subdirs = [p for p in tenable_root.iterdir() if p.is_dir()]
    if len(subdirs) == 1:
        only = subdirs[0]
        if any(looks_like_tenable_category_dir(p) for p in only.iterdir()):
            return only

    return tenable_root


def list_tenable_tag_dirs(tenable_root: Path) -> Dict[Tuple[str, str], Path]:
    out: Dict[Tuple[str, str], Path] = {}

    if not tenable_root.is_dir():
        return out

    real_root = resolve_tenable_base(tenable_root)

    for category_dir in real_root.iterdir():
        if not category_dir.is_dir():
            continue

        if not looks_like_tenable_category_dir(category_dir):
            continue

        for value_dir in category_dir.iterdir():
            if value_dir.is_dir():
                out[(norm(category_dir.name), norm(value_dir.name))] = value_dir

    return out


def resolve_amp_dir(target: Target, amp_dirs: Dict[str, Path]) -> Optional[Path]:
    return amp_dirs.get(norm(target.amp_name))


def resolve_tenable_dir(target: Target, tenable_dirs: Dict[Tuple[str, str], Path]) -> Optional[Path]:
    return tenable_dirs.get((norm(target.tenable_category), norm(target.tenable_value)))


def collect_images_for_target(
    target: Target,
    amp_dir: Optional[Path],
    tenable_dir: Optional[Path],
) -> Tuple[Dict[str, bytes], List[str]]:
    replacements: Dict[str, bytes] = {}
    warnings: List[str] = []

    amp_missing_reported = False
    tenable_missing_reported = False

    for media_path, (source, filename) in DOCX_IMAGE_MAP.items():
        root = amp_dir if source == "amp" else tenable_dir

        if root is None:
            if source == "amp" and not amp_missing_reported:
                warnings.append(
                    f"{target.display_name}: no se encontró carpeta AMP exacta '{target.amp_name}'. Se conservarán imágenes AMP originales."
                )
                amp_missing_reported = True
            elif source == "tenable" and not tenable_missing_reported:
                warnings.append(
                    f"{target.display_name}: no se encontró carpeta Tenable exacta '{target.tenable_category} -> {target.tenable_value}'. Se conservarán imágenes Tenable originales."
                )
                tenable_missing_reported = True
            continue

        file_path = root / filename
        if not file_path.is_file():
            warnings.append(
                f"{target.display_name}: falta {source}/{filename} en '{root}'. Se conserva imagen original."
            )
            continue

        replacements[media_path] = file_path.read_bytes()

    return replacements, warnings


def replace_docx_images(template_path: Path, output_path: Path, replacements: Dict[str, bytes]) -> None:
    with zipfile.ZipFile(template_path, "r") as zin, zipfile.ZipFile(
        output_path, "w", compression=zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename in replacements:
                data = replacements[item.filename]
            zout.writestr(item, data)


def replace_text_in_docx(
    docx_path: Path,
    old_entity: str,
    new_entity: str,
    old_period: Optional[str],
    new_period: Optional[str],
) -> None:
    try:
        from docx import Document  # type: ignore
    except Exception as exc:
        raise RuntimeError("python-docx no está disponible para reemplazar texto.") from exc

    def repl(text: str) -> str:
        if not text:
            return text
        text = text.replace(old_entity, new_entity)
        if old_period and new_period:
            text = text.replace(old_period, new_period)
        return text

    doc = Document(str(docx_path))

    def process_paragraphs(paragraphs):
        for p in paragraphs:
            for run in p.runs:
                run.text = repl(run.text)

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        process_paragraphs(section.header.paragraphs)
        process_paragraphs(section.footer.paragraphs)

        for table in section.header.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

        for table in section.footer.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)

    doc.save(str(docx_path))


def build_report_for_target(
    target: Target,
    template_path: Path,
    output_dir: Path,
    amp_dirs: Dict[str, Path],
    tenable_dirs: Dict[Tuple[str, str], Path],
) -> Tuple[Optional[Path], List[str]]:
    amp_dir = resolve_amp_dir(target, amp_dirs)
    tenable_dir = resolve_tenable_dir(target, tenable_dirs)

    replacements, warnings = collect_images_for_target(target, amp_dir, tenable_dir)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_file = output_dir / output_filename_for(target)
    replace_docx_images(template_path, out_file, replacements)

    if REPLACE_TEXT_INSIDE_DOCX:
        replace_text_in_docx(
            out_file,
            old_entity=TEMPLATE_ENTITY_NAME,
            new_entity=target.display_name,
            old_period=TEMPLATE_PERIOD_TEXT,
            new_period=REPORT_PERIOD or None,
        )

    return out_file, warnings


def filter_targets(targets: Sequence[Target], only: Optional[Sequence[str]]) -> List[Target]:
    if not only:
        return list(targets)

    wanted = {norm(x) for x in only}
    out: List[Target] = []

    for t in targets:
        names = {
            norm(t.display_name),
            norm(t.amp_name),
            norm(t.tenable_value),
            norm(t.output_filename or ""),
        }
        if names & wanted:
            out.append(t)

    return out


def dedupe_keep_order(items: List[str]) -> List[str]:
    seen = set()
    out = []
    for item in items:
        if item not in seen:
            seen.add(item)
            out.append(item)
    return out


# ============================================================
# MAIN
# ============================================================
def main(argv: Optional[Sequence[str]] = None) -> int:
    parser = argparse.ArgumentParser(
        description="Genera imágenes AMP/Tenable, arma DOCX y limpia temporales."
    )
    parser.add_argument("--solo", nargs="+", help="Procesa únicamente estas entidades.")
    args = parser.parse_args(argv)

    require_exists(AMP_SCRIPT, "AMP_SCRIPT")
    require_exists(TENABLE_SCRIPT, "TENABLE_SCRIPT")
    require_exists(TEMPLATE_DOCX, "TEMPLATE_DOCX")

    selected_targets = filter_targets(TARGETS, args.solo)
    if not selected_targets:
        print("[ERR] No hubo coincidencias en TARGETS para --solo.")
        return 2

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    with tempfile.TemporaryDirectory(prefix="reportes_amp_tenable_") as tmp:
        tmp_path = Path(tmp)
        amp_workdir = tmp_path / "amp_runner"
        tenable_workdir = tmp_path / "tenable_runner"
        amp_workdir.mkdir(parents=True, exist_ok=True)
        tenable_workdir.mkdir(parents=True, exist_ok=True)

        print(f"[INFO] Temporales en: {tmp_path}")

        print("[INFO] Generando imágenes de AMP...")
        run_python_script(AMP_SCRIPT, amp_workdir)
        amp_root = amp_workdir / "amp_reportes_unificado"
        require_exists(amp_root, "salida AMP")

        print("[INFO] Generando imágenes de Tenable...")
        tenable_root = tenable_workdir / "tenable_reportes_unificados"
        run_python_script(
            TENABLE_SCRIPT,
            tenable_workdir,
            extra_env={"TENABLE_REPORTS_OUT_DIR": str(tenable_root)},
        )
        require_exists(tenable_root, "salida Tenable")

        amp_dirs = list_amp_group_dirs(amp_root)
        tenable_dirs = list_tenable_tag_dirs(tenable_root)

        print(f"[INFO] Carpetas AMP detectadas: {len(amp_dirs)}")
        print(f"[INFO] Carpetas Tenable detectadas: {len(tenable_dirs)}")

        generated = 0
        all_warnings: List[str] = []

        for target in selected_targets:
            print(f"[INFO] Generando DOCX para: {target.display_name}")
            out_file, warnings = build_report_for_target(
                target,
                TEMPLATE_DOCX,
                OUTPUT_DIR,
                amp_dirs,
                tenable_dirs,
            )
            all_warnings.extend(warnings)

            if out_file:
                generated += 1
                print(f"[OK] {target.display_name} -> {out_file}")

        print("[INFO] Limpiando imágenes/carpetas temporales...")

    print(f"[DONE] DOCX generados: {generated}")
    print(f"[DONE] Carpeta final: {OUTPUT_DIR}")

    all_warnings = dedupe_keep_order(all_warnings)
    if all_warnings:
        print("\n[WARN] Resumen de advertencias:")
        for w in all_warnings:
            print(f" - {w}")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())